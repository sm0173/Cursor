# -*- coding: utf-8 -*-
"""Экспорт Отчет_этап4.md в Отчет_этап4.docx (базовый разбор Markdown)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def _add_runs_with_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def _is_separator_row(line: str) -> bool:
    s = line.strip().strip("|").replace(" ", "")
    return bool(re.match(r"^:?-+:?(?:\|:?-+:?)*$", s))


def md_to_docx(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if _is_table_row(stripped):
            rows: list[list[str]] = []
            header = _parse_table_row(stripped)
            rows.append(header)
            i += 1
            if i < len(lines) and _is_separator_row(lines[i].strip()):
                i += 1
            while i < len(lines) and _is_table_row(lines[i].strip()):
                rows.append(_parse_table_row(lines[i].strip()))
                i += 1
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell_text = row[ci] if ci < len(row) else ""
                    table.rows[ri].cells[ci].text = cell_text
            doc.add_paragraph()
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            body = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, body)
            i += 1
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_runs_with_bold(p, stripped)
        i += 1


def main() -> None:
    root = Path(__file__).resolve().parent
    md_path = root / "Отчет_этап4.md"
    out_path = root / "Отчет_этап4.docx"
    doc = Document()
    _set_styles(doc)
    md_to_docx(md_path.read_text(encoding="utf-8"), doc)
    try:
        doc.save(out_path)
        print("OK", out_path)
    except PermissionError:
        alt = root / "Отчет_этап4_расширенный.docx"
        doc.save(alt)
        print("WARN: основной .docx занят, сохранено как", alt)


if __name__ == "__main__":
    main()
